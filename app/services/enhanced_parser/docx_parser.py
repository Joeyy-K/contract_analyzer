import io
from typing import Dict, Any, Optional
import zipfile

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class DOCXParser:
    """
    Enhanced DOCX parser that extracts text, metadata, and structural elements.
    """
    
    def __init__(self):
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file and extract text with metadata."""
        result = {
            'text': '',
            'metadata': {},
            'success': False,
            'errors': []
        }
        
        try:
            doc = Document(file_path)
            result.update(self._extract_content(doc))
            result['success'] = True
        except Exception as e:
            result['errors'].append(f"Error parsing DOCX: {str(e)}")
            result['success'] = False
        
        return result
    
    def parse_bytes(self, file_content: bytes) -> Dict[str, Any]:
        """Parse DOCX from bytes content."""
        result = {
            'text': '',
            'metadata': {},
            'success': False,
            'errors': []
        }
        
        try:
            doc = Document(io.BytesIO(file_content))
            result.update(self._extract_content(doc))
            result['success'] = True
        except Exception as e:
            result['errors'].append(f"Error parsing DOCX: {str(e)}")
            result['success'] = False
        
        return result
    
    def _extract_content(self, doc: Document) -> Dict[str, Any]:
        """Extract all content from DOCX document."""
        content = {
            'text': '',
            'metadata': {}
        }
        
        # Extract metadata
        content['metadata'] = self._extract_metadata(doc)
        
        # Extract main document text
        text_parts = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    text_parts.append(f"\n{'#' * int(para.style.name[-1])} {para.text}\n")
                else:
                    text_parts.append(para.text)
        
        # Process tables
        for table_num, table in enumerate(doc.tables):
            table_text = f"\n--- Table {table_num + 1} ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text += row_text + "\n"
            text_parts.append(table_text)
        
        # Extract headers and footers
        headers_footers = self._extract_headers_footers(doc)
        if headers_footers:
            text_parts.append(f"\n--- Headers/Footers ---\n{headers_footers}")
        
        content['text'] = "\n".join(text_parts)
        return content
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {}
        
        try:
            # Core properties
            core_props = doc.core_properties
            metadata.update({
                'title': getattr(core_props, 'title', '') or '',
                'author': getattr(core_props, 'author', '') or '',
                'subject': getattr(core_props, 'subject', '') or '',
                'created': str(getattr(core_props, 'created', '') or ''),
                'modified': str(getattr(core_props, 'modified', '') or ''),
                'last_modified_by': getattr(core_props, 'last_modified_by', '') or '',
                'revision': str(getattr(core_props, 'revision', '') or ''),
                'category': getattr(core_props, 'category', '') or '',
                'comments': getattr(core_props, 'comments', '') or ''
            })
            
            # Document statistics
            metadata.update({
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            })
            
        except Exception as e:
            metadata['metadata_error'] = str(e)
        
        return metadata
    
    def _extract_headers_footers(self, doc: Document) -> str:
        """Extract text from headers and footers."""
        headers_footers = []
        
        try:
            for section in doc.sections:
                # Extract header
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            headers_footers.append(f"Header: {para.text}")
                
                # Extract footer
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            headers_footers.append(f"Footer: {para.text}")
                            
        except Exception:
            pass  # Headers/footers extraction is optional
        
        return "\n".join(headers_footers)